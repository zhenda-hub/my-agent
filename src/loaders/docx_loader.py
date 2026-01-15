"""Word 文档加载器"""
from pathlib import Path
from typing import List
from docx import Document as DocxDocument
from src.loaders.base import BaseLoader, Document


class DocxLoader(BaseLoader):
    """Word (.docx) 文档加载器"""

    def load(self, path: str) -> List[Document]:
        """
        加载 Word 文档

        Args:
            path: Word 文件路径

        Returns:
            文档列表（整个文档作为一个文档）
        """
        path_obj = Path(path)
        if not path_obj.exists():
            raise FileNotFoundError(f"Word file not found: {path}")

        doc = DocxDocument(path)
        paragraphs = []

        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                paragraphs.append(paragraph.text)

        content = "\n".join(paragraphs)

        return [
            Document(
                content=content,
                metadata={
                    "type": "docx",
                    "paragraphs_count": len(doc.paragraphs),
                },
                source=str(path_obj),
            )
        ]
